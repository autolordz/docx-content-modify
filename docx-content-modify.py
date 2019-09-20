# -*- coding: utf-8 -*-
# Copyright (c) 2018 Autoz https://github.com/autolordz

# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:

# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.

# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.

"""
Created on Wed Sep 11 17:33:29 2019

@author: autol
"""

#%%
import os,re
import pandas as pd
from docx import Document
#import sys
#sys.stderr = open(os.devnull, "w")  # silence stderr
from globalvar import *

#%%
import util as ut
from df_progress import df_fill_infos,df_oa_append,df_read_fix,merge_group_cases
from df_transform import df_check_format
from df_transform import make_adr,make_agent,merge_user,merge_usr_agent_adr
from df_transform import reclean_data,sort_data

#%%
print('''
Postal Notes Automatically Generate App

Updated on Thu Sept 19 2019

Depends on: python-docx,pandas,StyleFrame,configparser

@author: Autoz (autolordz@gmail.com)
''')

#%%

print_log('''>>> 正在处理...
    主表路径 = %s
    指定案件 = %s
    指定日期 = %s
    指定条数 = %s
    '''%(os.path.abspath(data_xlsx),
        data_case_codes,
        data_date_range,
        data_last_lines,
        )
    )

if not os.path.exists(data_xlsx):
    ut.save_adjust_xlsx(pd.DataFrame(columns=titles_main),data_xlsx,width=40)
    print_log('>>> %s 文件不存在...重新生成'%(data_xlsx))

#%%
df = pd.read_excel(data_xlsx,sort=False).fillna('') #真正读取记录位置
df = df_read_fix(df);df # fix empty data columns
df,ocodes = df_oa_append(df) # append oa data and sava # 合并前记录案号
df = merge_group_cases(df);df # merge group and save
df = df[df['原一审案号'].isin(ocodes)];df # 合并后找回记录案号
df = df_fill_infos(df) # 填充判决书内容 # filled and save
df = ut.titles_trans_columns(df,titles_cn);df # 中译英方便后面处理
if flag_check_postal:
    df.apply(lambda x:df_check_format(x), axis=1)

if 0<len(df)<10:
    print_log('>>> 将要打印【%s条】=> %s '%(len(df),
                                       df['number'].to_list()))

#%% df tramsfrom stream 数据转换流程

if len(df) and flag_to_postal:
    try:
        print_log('\n>>> 开始生成新数据 data_main_temp... ')
        '''获取 datetime|number'''
        number = df[titles_en[:2]]
        number = number.reset_index()
        number.columns.values[0] = 'level_0'
        # user
        '''获取所有用户名包括曾用名'''
        user = df['uname']
        user = user[user != '']
        user = user.str.strip().str.split(r'[,，。]',expand=True).stack() # divide user
        user = user.str.strip().str.split(r'[:]',expand=True)# divide character
        user = user[1].str.strip().str.split(r'[、]',expand=True).stack().to_frame(name = 'uname')
        user = user.reset_index().drop(['level_1','level_2'],axis=1)
        # agent and address
        agent_adr = df[['aname','address']]
        opt = agent_adr.any()
        agent = df['aname']
        adr = df['address']

        if all(opt):
            print_log('>>> 有【代理人】和【地址】...正在处理...')
            adr = make_adr(adr,fix_aname=user['uname'].tolist())
            agent = make_agent(agent,fix_aname=adr['clean_aname'].tolist()) #获取代理人
            usr_agent = merge_user(user,agent)
            df_x = reclean_data(merge_usr_agent_adr(usr_agent,adr))
            df_x = sort_data(df_x,number)
        elif opt.address:
            print_log('>>> 只有【地址】...正在处理...')
            adr = make_adr(adr,fix_aname=user['uname'].tolist())
            adr['uname'] = adr['clean_aname']
            adr = merge_user(user,adr)
            adr = adr.assign(aname='')
            df_x = reclean_data(adr)
            df_x = sort_data(df_x,number)
        elif opt.aname:
            print_log('>>> 只有【代理人】...正在处理...')
            agent = make_agent(agent)
            agent = merge_user(user,agent)
            agent = agent.assign(address='')
            df_x = reclean_data(agent)
            df_x = sort_data(df_x,number)
        else:
            print_log('>>> 缺失【代理人】和【地址】...正在处理...')
            agent_adr.index.name = 'level_0'
            agent_adr.reset_index(inplace=True)
            df_x = pd.merge(user,agent_adr,how='left',on=['level_0']).fillna('')
            df_x = sort_data(df_x,number)

        if len(df_x):
            data_tmp = os.path.splitext(data_xlsx)[0]+"_tmp.xlsx"
            df_save = df_x.copy()
            df_save.columns = ut.titles_switch(df_save.columns.tolist())
            df_save = ut.save_adjust_xlsx(df_save,data_tmp,width=40)

    except Exception as e:
        input_exit('>>> 错误 \'%s\' 生成数据失败,请检查源 \'%s\' 文件...退出...'%(e,data_xlsx))

#%% print postal sheets 打印邮单流程

def re_write_text(x):
    '''re-write postal sheet content from df rows'''

    doc = Document(sheet_docx)
    doc.styles['Normal'].font.bold = True
    uname = str(x['uname']);aname = str(x['aname'])
    agent_text = aname if aname else uname
    user_text = '' if uname in agent_text else '代 '+ uname
    number_text = str(x['number'])
    address_text = str(x['address'])

    # 以下填充均对于模板sheet.doc
    try:
        para = doc.paragraphs[9]  # No.9 line is agent name
        text = re.sub(r'[\w（）()]+',agent_text,para.text)
        para.clear().add_run(text)

        para = doc.paragraphs[11]  # No.11 line is user name
        text = re.sub(r'代 \w+',user_text,para.text)
        para.clear().add_run(text)


        para = doc.paragraphs[13]  # No.13 line is number and address
        text = re.sub(ut.path_code_ix,number_text,para.text)
        para.clear().add_run(text)
        text = re.sub(r'(?<=\s)\w+市.*',address_text,para.text)
        para.clear().add_run(text)
    except Exception as e:
        print_log('错误 \'%s\' 替换文本 => \'%s\' 失败！！！' %(e,para.text))

    sheet_file = number_text+'_'+agent_text+'_'+user_text+'_'+address_text+'.docx'
    sheet_file = re.sub(r'[\/\\\:\*\?\"\<\>]',' ',sheet_file) # keep rename legal

    if os.path.exists(ut.parse_subpath(postal_path,sheet_file)):
        if ut.flag_check_postal:print_log('>>> 邮单已存在！！！ <= %s'%sheet_file)
        return ''

    if not agent_text:
        if flag_check_postal:print_log('>>> 【代理人】暂缺！！！ <= %s'%sheet_file)
        return ''

    if not address_text:
        if flag_check_postal:print_log('>>> 【地址】暂缺！！！ <= %s'%sheet_file)
        return ''
    try:
        doc.save(ut.parse_subpath(postal_path,sheet_file))
        print_log('>>> 已生成邮单 => %s'%sheet_file)
        return sheet_file
    except Exception as e:
        input_exit('>>> 生成失败！！！ => %s ...任意键退出'%e)
    return ''


if len(df) and flag_to_postal:
    print_log('\n>>> 正在输出邮单...\n')
    if not os.path.exists(sheet_docx):
        input_exit('>>> 没有找到邮单模板 %s...任意键退出'%sheet_docx)
    df_p = df_x.apply(re_write_text,axis = 1)
    count = len(df_p[df_p != ''])
    codes = df_x['number'].astype(str)
    dates = df_x['datetime'].astype(str)
    codesrange = codes.iloc[0] if codes.iloc[0] == codes.iloc[-1] else ('%s:%s'%(codes.iloc[0],codes.iloc[-1]))
    datesrange = dates.iloc[0] if dates.iloc[0] == dates.iloc[-1] else ('%s:%s'%(dates.iloc[0],dates.iloc[-1]))
    print_log('\n>>> 最终生成邮单【%s条】范围: 【%s】日期:【%s】'%(count,codesrange,datesrange))

    del df_x,df_p,codes,dates
    del user,number,agent,adr,df,agent_adr,opt

input_exit('>>> 全部完成,可以回顾记录...任意键退出')
