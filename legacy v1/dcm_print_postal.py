# -*- coding: utf-8 -*-
"""
Created on Thu Dec 12 17:09:28 2019

@author: autol
"""

import re
from docx import Document
import dcm_util as ut
from dcm_globalvar import *
locals().update(var.to_dict()) # 设置读取的全局变量


#%% print postal sheets 打印邮单流程

def fill_postal_save(x):
    '''按df行填充并输出邮单文件'''

    doc = Document(sheet_docx)
    doc.styles['Normal'].font.bold = True
    uname = str(x['uname']);aname = str(x['aname'])
    agent_text = aname if aname else uname
    user_text = '' if uname in agent_text else '代 '+ uname
    number_text = str(x['number'])
    address_text = str(x['address'])
    
    print("打印信息...用户【%s】...代理人【%s】...地址【%s】..."%(user_text,agent_text,address_text))

    # 以下填充均对于模板sheet.doc，假如模板位置有变，这里需要修改
    try:
        para = doc.paragraphs[9]  # No.9 line is agent name
        text = re.sub(r'[\w（）()]+',agent_text,para.text)
        para.clear().add_run(text)

        para = doc.paragraphs[11]  # No.11 line is user name
        text = re.sub(r'代 \w+',user_text,para.text)
        para.clear().add_run(text)

        para = doc.paragraphs[13]  # No.13 line is number and address
        text = re.sub(ut.path_code_ix,number_text,para.text)
        para.clear().add_run(text)
        text = re.sub(r'(?<=\s)\w+市.*',address_text,para.text)
        para.clear().add_run(text)
    except Exception as e:
        print_log('错误 \'%s\' 替换文本 => \'%s\' 失败！！！' %(e,para.text))

    sheet_file = number_text+'_'+agent_text+'_'+user_text+'_'+address_text+'.docx'
    sheet_file = re.sub(r'[\/\\\:\*\?\"\<\>]',' ',sheet_file) # keep rename legal

    if os.path.exists(ut.parse_subpath(postal_path,sheet_file)):
        if ut.flag_check_postal:print_log('>>> 邮单已存在！！！ <= %s'%sheet_file)
        return ''

    if not agent_text:
        if flag_check_postal:print_log('>>> 【代理人】暂缺！！！ <= %s'%sheet_file)
        return ''
#
#    if not address_text:
#        if flag_check_postal:print_log('>>> 【地址】暂缺！！！ <= %s'%sheet_file)
#        return ''
    try:
        doc.save(ut.parse_subpath(postal_path,sheet_file))
        print_log('>>> 已生成邮单 => %s'%sheet_file)
        return sheet_file + ' Yes!!'
    except Exception as e:
        input_exit('>>> 生成失败！！！ => %s ...任意键退出'%e)
    return ''


