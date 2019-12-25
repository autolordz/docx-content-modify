# -*- coding: utf-8 -*-
"""
Created on Wed Sep 11 11:44:28 2019

@author: autol
"""

import os,re
from glob import glob
import pandas as pd
from docx import Document
#%%
from dcm_util import check_codes,case_codes_fix,split_list,parse_subpath
from dcm_globalvar import *

#%% 读取判决书jdocs代码

def read_jdocs_table(tables):
    codes = ''
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    x = paragraph.text
                    if re.search(path_code_ix,x) and len(x) < 25:
                        codes = case_codes_fix(x)
                        break
    return codes

def get_jdocs_infos(doc,lines = 20):# search at least 20 lines docs
    '''get pre address from judgment docs, return docs pre code and address'''
    adrs = {};codes = ''
    try:tables = Document(doc).tables
    except Exception as e:
        print_log('读取错误 %s ,docx文档问题,请重新另存为,或关闭已打开的docx文档'%e)
        return codes,adrs
    if tables: codes = read_jdocs_table(tables)
    paras = Document(doc).paragraphs
    if not paras:
        return codes,adrs
    if len(paras) > 20: # 多于20行就扫描一半内容
        lines = int(len(paras)/2)
    parass = paras[:lines]
    for i,para in enumerate(parass):
        x = para.text.strip()
        if len(x) > 150: continue # 段落大于150字就跳过
        if re.search(path_code_ix,x) and len(x) < 25:
            codes = case_codes_fix(x);continue # codes
        cond3 = re.search(r'法定代表|诉讼|代理人|判决|律师|请求|证据|辩称|辩论|不服',x) # 跳过非人员信息
        cond4 = re.search(r'上市|省略|区别|借款|保证|签订',x) # 跳过非人员信息,模糊
        cond1 = re.search(r'(?<=[：:]).*?(?=[,，。])',x) # 通过间隔提取
        cond2 = re.search(r'.*?[省市州县区乡镇村]',x) # 地址规则
        if cond3:continue
        if cond4:continue
        if cond1 and cond2:
            '''
            Todo: get user and address
            Usage: '被上诉人（原审被告）：张三，男，1977年7月7日出生，汉族，住XX自治区(省)XX市XX区1212。现住XX省XX市XX区3434'
            -> {'张三': 'XX省XX市XX区3434'}
            '''
            try:
                name = re.search(r'(?<=[：:]).*?(?=[,，。])|$',x).group(0).strip()
                name = re.sub(r'[(（][下称|原名|反诉|变更前].*?[）)]','',name) # filter some special names,notice here will add some words for filter
                z = split_list(r'[,，:：.。]',x)
                z = [re.sub(r'户[籍口]|居住|身份证|所在地|住所地?|住址?|^[现原]住?','',y) for y in z if re.search(r'.*?[省市州县区乡镇村]',y)][-1] # 几个地址选最后一个 remain only address
                adr = {name:''.join(z)}
                adrs.update(adr)
            except Exception as e:
                print_log('获取信息失败 =>',e)
    return codes,adrs

def rename_jdoc_x(doc,codes):
    '''rename only judgment doc files'''
    jdoc_name = os.path.join(os.path.split(doc)[0],'判决书_'+codes+'.docx')
    if not codes in doc:# os.path.exists(jdoc_name)
        try:
            os.rename(doc,jdoc_name)
            return 1
        except Exception as e:
            print_log(e)
            os.remove(doc)
            return 0
    return 0

def get_all_jdocs(docs):
    '''主要获取的入口'''
    numlist=[]; nadr = []
    for doc in docs:
        codes,adrs = get_jdocs_infos(doc)
        if codes:
            rename_jdoc_x(doc,codes)
        numlist.append(codes)
        nadr.append(adrs)
        if flag_check_jdocs and codes:
            print_log('>>> 判决书信息 【%s】-【%s人】-%s \n'%(codes,len(adrs),adrs))
    numlist = list(map(case_codes_fix,numlist))
    return pd.DataFrame({'判决书源号':numlist,'new_adr':nadr})


def rename_jdocs_codes_x(d,r,old_codes):
    '''add jdoc current case codes for reference 判决书改名，包括源案号'''
    if str(r[old_codes]) in str(d):
        nd = os.path.join(os.path.split(d)[0],'判决书_'+str(r['案号']) +'_原_'+ str(r[old_codes]) + '.docx')
        if(d == nd): # 相同则返回
            return d
        try: # 不同则命名，检测源文件存在
            if os.path.exists(nd):
                os.remove(nd)
            os.rename(d,nd)
            print('>>> 重命名判决书 => ',nd)
        except Exception as e:
            print_log(e)
        return nd
    return d

def rename_jdocs_codes(dfo):
    '''rename jdocs with new codes'''
    old_codes='判决书源号'
    docs = glob(parse_subpath(jdocs_path,'判决书_*.docx'))
    df = dfo[dfo[old_codes] != '']
    if docs:
        for doc in docs:
            for (i,dfr) in df.iterrows():
                if check_codes(dfr[old_codes]) and str(dfr[old_codes]) in doc:
                    rename_jdocs_codes_x(doc,dfr,old_codes)
                    break
    return None