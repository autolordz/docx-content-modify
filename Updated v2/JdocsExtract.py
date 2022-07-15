# -*- coding: utf-8 -*-
"""
Created on Mon Jun 20 17:41:36 2022

@author: Autozhz
"""

import os,re,glob,sys,time,random,string
import pandas as pd
from docx import Document
from CommomClass import CommomClass

class JdocsExtract(CommomClass):
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
    
    def check_table_content(self,doc_name):
        codes = '';isanhao = 0
        try:
            tables = Document(doc_name).tables
            if tables: # read if there's table in jdocs, find codes in table very first 
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                match_anhao = re.search(r'[(（]\d+[）)].*\d+(?=号)',p.text) # codes limited with 25 digits
                                if match_anhao:
                                    codes = match_anhao.group(0)+'号'; isanhao=1; break
        except Exception as e:
            err = '读取文件【%s】错误 ,word文档的问题，或格式docx和doc混淆，或文档已打开，或文档损坏 Skip It.\n %s'%(doc_name,e)
            print(err)
        return codes,isanhao
    
    def get_member_info_all(self,doc_folder,**kwargs):
        print('\n'+'Start Reading All 判决书'.center(30,'*'))
        jdoc_files = glob.glob(doc_folder+'\\*.docx')
        df = self.read_file(**kwargs)
        for doc_name in jdoc_files:
            df1 = self.get_member_info(doc_name)
            if not df1 is None:
                df = pd.concat([df,df1])

        df = self.sort_duplicates_data(df,sortlist=['原审案号','当事人'])
        self.save_file(df_output=df,
                       **kwargs)
        print('END'.center(30,'*') + '\n')
        return df
    
    def get_member_info(self,doc_name):
        try:
            dd = Document(doc_name).paragraphs
        except Exception as e:
            print(e); return
        tlist = []
        anhao,isanhao = self.check_table_content(doc_name)
        anhao = '';isanhao = 0
        for p in dd:
            if not isanhao and len(p.text) < 50:
                match_anhao = re.search(r'[(（]\d+[）)].*\d+(?=号)',p.text)
                if match_anhao:
                    anhao = (match_anhao.group(0)+'号').replace(r'(',r'（').replace(r')',r'）')
                    isanhao = 1; continue
            if len(p.text) > 150: continue
            role_rgx = '原告|被告|原审原告|原审被告|第三人|原审第三人|申请人|被申请人|上诉人|被上诉人|异议人|执行人'
            match_member = re.search(r'%s'%role_rgx,p.text) # 找人员 # rest 负责人|经营者|法定代表人|法人|委托诉讼代理人|
            match_agent = re.search(r'(委托|诉讼)?代理人',p.text)
            if match_member and not match_agent:
                if match_member.span()[0]>3: continue # 角色不在开头位置，属于文章内容
                if re.search(r'(%s).*(证据|请求|称|意见)\w*(?=[：:])'%role_rgx,p.text) : continue # XX称，属于文章内容
                match_name = re.search('(?<=[：:])[^。，,]+(?=[。，,])',p.text) # find the member name 
                if match_name: 
                    role_title = re.search(r'[^：:]*(?=[：:])',p.text).group(0) #截取头 上诉人（一审被告）
                    if len(role_title) > 18: continue # 角色超长，属于文章内容
                    member = match_name.group(0).strip()
                    # if not re.search(r'合作社|酒店|公司',name) and len(name) > 5: continue # 去掉非公司人名和长度超过5的段落
                    # member = re.sub(r'[(（][下称|原名|反诉|变更前].*?[）)]','',member) # filter member name of some special members,notice here will add some words for filter
                    adrs_regex = r'^户[籍口]|居住|身份证|所在地|住所地?|住址?|[现原]?住|为?'
                    adrs = [re.sub(adrs_regex,'',y) \
                            for y in self.split_list(r'[,，:：.。]',p.text) \
                                if re.search(r'(%s).*?[省市州县区乡镇村]'%adrs_regex,y)] # 几个地址选最后一个 remain only address
                    adrs = adrs[-1] if len(adrs) > 0 else ''
                    df_row = {'原审案号':anhao,'角色':role_title,'当事人':member,'地址':adrs}
                    tlist.append(df_row) #直接获取信息
        if not tlist:
            print('请检查 判决书 %s 里"角色"是否没有冒号'%anhao)
        self.rename_jdocs(anhao, doc_name)
        df = pd.DataFrame(tlist)
        return df
    
    def rename_jdocs(self,anhao,doc_name): # for 的子内容
        # nd = os.path.join(os.path.split(d)[0],'判决书_'+anhao.strip() +'_原_'+ str(r[old_codes]) + '.docx')
        if '没有案号' in doc_name or '重复' in doc_name: return 0
        if not anhao:
            # t1 = time.ctime(os.path.getmtime(doc_name))
            # t2 = time.strftime('%Y%m%d%H%M%S',time.strptime(t1))
            t2 = ''.join(random.choice(string.ascii_letters) for _ in range(6))
            nd = '%s\\判决书_没有案号_%s.docx'%(os.path.split(doc_name)[0],t2)
        else:
            nd = '%s\\判决书_%s.docx'%(os.path.split(doc_name)[0],anhao.strip())
        if(doc_name == nd): # 相同则返回
            return 0
        try: # 不同则命名，检测源文件存在
            if os.path.exists(nd):
                postfix = ''.join(random.choice(string.ascii_letters) for _ in range(6))
                nd = nd.replace('.docx','_重复_%s.docx'%postfix)
            os.rename(doc_name,nd)
            print('>>> 重命名判决书 => ',nd)
        except Exception as e:
            print(e); return 0
        return 1

# from Global import *
# JdocsExtract1 = JdocsExtract()
# df = JdocsExtract1.get_member_info_all(doc_folder,
#                                        df_input_name=jdocs_extract_path,
#                                        df_output_name=jdocs_extract_path,
#                                        isSave=1)




