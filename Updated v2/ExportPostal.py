# -*- coding: utf-8 -*-
"""
Created on Sat Jun 18 16:51:29 2022

@author: Autozhz
"""

from docx import Document
import os,re,sys
from PrintLog import PrintLog
import pandas as pd
from CommomClass import CommomClass

class ExportPostal(CommomClass):
    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
    
    def generate_postal(self,**kwargs):
        self.postal_path = kwargs.get('postal_path',None)
        self.postal_template = kwargs.get('postal_template',None)
        
        count = 0
        df = self.read_file(**kwargs)
        for i,row in df.iterrows():
            count = count + self.fill_postal(row)
        
        self.pop_msg('成功生成 【%s】 张邮单'%count)
        self.open_path(self.postal_path, **kwargs)
        
        return df
    
    def fill_postal(self,x):
        
        postal_path = self.postal_path
        
        doc = Document(self.postal_template)
        doc.styles['Normal'].font.bold = True
        
        uname = str(x['当事人'])
        aname = re.search(r'[^,.，、；]*(?=[,.，、；])|\w+|$',
                          str(x['代理人'])).group(0) #多个名字只要一个
        tel = re.search(r'[^,.，、；]*(?=[,.，、；])|\w+|$',
                        str(x['联系'])).group(0) #多个tel只要一个
        case_number = str(x['案号'])
        address = str(x['地址'])
        judge = str(x['承办法官'])
        lawfirm = str(x['律所'])
        
        if not (uname and address):
            PrintLog.log('>>> 【%s】信息不全不打印！ '%case_number)
            return 0
        
        text_agent = aname if aname and aname != uname else uname
        text_user = '（代 %s）'%uname if aname and aname != uname else '' # 加入没agent则user不填,user变代理人
        
        # 以下填充均对于模板sheet.doc，假如模板位置有变，这里需要修改
        # 8 aname tel
        # 9 uname
        # 11 lawfirm
        # 12 case 
        # 13 address
        # 19 shujiyuan faguan
        try:
            para = doc.paragraphs[8] # find aname until tel
            text = re.sub(r'\S+',text_agent,para.text) # [^\s]+(?=\s)
            para.clear().add_run(text)
            
            para = doc.paragraphs[8] # tel
            text = re.sub(r'\S+$',tel,para.text) # (?<=\s)\d+
            para.clear().add_run(text)
        
            para = doc.paragraphs[9]  
            text = re.sub(r'[^\s].*[^\s]',text_user,para.text)
            para.clear().add_run(text)
            
            para = doc.paragraphs[11]  
            text = re.sub(r'\S+$',lawfirm,para.text)
            para.clear().add_run(text)
        
            para = doc.paragraphs[12]
            text = re.sub(r'[^\s].*[^\s]',case_number,para.text)
            para.clear().add_run(text)
            
            para = doc.paragraphs[13]
            text = re.sub(r'[^\s].*[^\s]',address,para.text) 
            para.clear().add_run(text)
            
            para = doc.paragraphs[18]
            text = re.sub(r'法官',judge,para.text) 
            para.clear().add_run(text)
            
        except Exception as e:
            PrintLog.log('出错了 \'%s\' 替换文本失败！' %(e))
            return 0
        
        sheet_name = '%s_%s_%s的邮单.docx'%(case_number,text_agent,text_user)
        sheet_name = re.sub(r'[\/\\\:\*\?\"\<\>]',' ',sheet_name) # keep rename legal
        
        postal_full_path = os.path.join(postal_path,sheet_name)
        
        if os.path.exists(postal_full_path):
            PrintLog.log('>>> 邮单已存在！ <= %s'%sheet_name)
            return 0 

        try:
            doc.save(postal_full_path)
            PrintLog.log('>>> 已生成邮单 => %s'%sheet_name)
            return 1
        except Exception as e:
            PrintLog.log('出错了 \'%s\' 保存文件失败！' %(e))
        return 0

# from Global import *
# ExportPostal1 = ExportPostal()
# df = ExportPostal1.generate_postal(
#                             df_input_name = df_expand_path,
#                             isOpenPath = 1,
#                             **kwargs)

#%%
# from docx import Document
# doc = Document('D:\\Python\\youdanji\\template.docx')
# doc.styles['Normal'].font.bold = True

# # 8 uname tel
# # 9 aname
# # 12 case
# # 13 address
# # 18 shujiyuan faguan

# for i,row in enumerate(doc.paragraphs):
#     print(i,row.text)
    
